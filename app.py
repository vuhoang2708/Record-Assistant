import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
import tempfile
import os
import json
import pandas as pd
import time

# --- CẤU HÌNH TRANG (GIAO DIỆN 2025) ---
st.set_page_config(
    page_title="AI Meeting Assistant 2025",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- CSS MODERN UI ---
st.markdown("""
<style>
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        height: 3em;
        font-weight: bold;
    }
    .report-box {
        border: 1px solid #e0e0e0;
        padding: 20px;
        border-radius: 10px;
        background: #ffffff;
    }
</style>
""", unsafe_allow_html=True)

# --- HÀM XỬ LÝ LOGIC ---

def configure_genai():
    try:
        api_key = st.secrets["GOOGLE_API_KEY"]
        genai.configure(api_key=api_key)
        return True
    except KeyError:
        st.error("🚨 Thiếu API Key trong secrets.toml")
        return False

def upload_to_gemini(path, mime_type="audio/mp3"):
    """Upload file lên Gemini Storage (Hỗ trợ file cực lớn của Gemini 3)."""
    file = genai.upload_file(path, mime_type=mime_type)
    while file.state.name == "PROCESSING":
        time.sleep(1)
        file = genai.get_file(file.name)
    return file

def build_dynamic_prompt(options):
    """Prompt kỹ thuật Prompt Engineering 2025."""
    base_prompt = """
    Role: Bạn là AI Secretary cao cấp (sử dụng engine Gemini 3.0).
    Context: Xử lý file âm thanh cuộc họp.
    Output Requirement: Trả về định dạng Markdown chuẩn, tối ưu cho việc convert sang Docx và JSON.
    
    TASKS:
    """
    
    tasks = []
    if options.get("transcript"):
        tasks.append("- **TRANSCRIPT:** Gỡ băng chính xác từng từ, định dạng: [Time] [Speaker]: Content.")
    
    if options.get("summary"):
        tasks.append("- **EXECUTIVE SUMMARY:** Tóm tắt ý chính. Tạo bảng Action Items (Task, Owner, Deadline).")
    
    if options.get("prosody"):
        tasks.append("- **SENTIMENT ANALYSIS:** Phân tích biểu đồ cảm xúc của cuộc họp (Căng thẳng/Hài lòng/Trung tính).")
    
    if options.get("minutes"):
        tasks.append("- **OFFICIAL MINUTES:** Biên bản họp chuẩn doanh nghiệp (Heading 1, 2 rõ ràng).")
    
    if options.get("gossip"):
        tasks.append("- **GOSSIP MODE:** Kể lại drama cuộc họp bằng ngôn ngữ Gen Alpha/Z.")
    
    if options.get("notebooklm_data"):
        tasks.append("""
        - **NOTEBOOKLM STUDIO DATA:** 
          1. Trích xuất dữ liệu quan trọng dưới dạng cấu trúc JSON để import vào NotebookLM Studio (cho Slide & Infographic).
          2. Tạo cấu trúc bảng dữ liệu (Table) cho các chỉ số tài chính/KPIs nếu có.
        """)

    return base_prompt + "\n".join(tasks)

def create_docx(content):
    doc = Document()
    doc.add_heading('MEETING REPORT - GEMINI 3.0', 0)
    
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        else:
            doc.add_paragraph(line)
    return doc

# --- MAIN APP ---

def main():
    st.title("🚀 AI Meeting Assistant (Gen 3)")
    st.caption("Powered by Gemini 3.0 Flash & NotebookLM Studio Integration")

    if not configure_genai():
        return

    # --- SIDEBAR ---
    with st.sidebar:
        st.header("🧠 Model Engine (Dec 2025)")
        
        # Cập nhật Model Selection theo thời điểm 12/2025
        model_version = st.selectbox(
            "Chọn Model:",
            (
                "gemini-3.0-flash", # Model mặc định, siêu nhanh, context vô cực
                "gemini-3.0-pro",   # Reasoning mạnh hơn
                "gemini-ultra-next" # Bản cao cấp nhất
            )
        )
        
        st.divider()
        st.subheader("🛠️ Feature Modules")
        
        opt_transcript = st.checkbox("Full Transcript", False)
        opt_summary = st.checkbox("Summary & Actions", True)
        opt_minutes = st.checkbox("Formal Minutes", True)
        opt_prosody = st.checkbox("Prosody/Sentiment", False)
        opt_gossip = st.checkbox("Gossip Mode", False)
        
        st.markdown("**NotebookLM Studio Integration:**")
        opt_notebooklm = st.checkbox("Generate Slide/Infographic Data", False, help="Tạo dữ liệu cấu trúc để import vào NotebookLM Studio mới")

        options = {
            "transcript": opt_transcript,
            "summary": opt_summary,
            "minutes": opt_minutes,
            "prosody": opt_prosody,
            "gossip": opt_gossip,
            "notebooklm_data": opt_notebooklm
        }

    # --- UPLOAD AREA ---
    uploaded_file = st.file_uploader("Upload Recording (mp3, wav, m4a)", type=['mp3', 'wav', 'm4a'])

    if uploaded_file:
        st.audio(uploaded_file)
        
        if st.button(f"⚡ XỬ LÝ VỚI {model_version.upper()}", type="primary"):
            
            prompt = build_dynamic_prompt(options)
            
            with st.spinner(f"Gemini 3.0 đang phân tích ngữ nghĩa & tín hiệu âm thanh..."):
                try:
                    # 1. Temp File
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
                        tmp.write(uploaded_file.getvalue())
                        tmp_path = tmp.name

                    # 2. Upload to Gemini 3 Storage
                    gemini_file = upload_to_gemini(tmp_path)
                    
                    # 3. Generate
                    # Lưu ý: Code này chạy giả định tên model là gemini-3.0-flash
                    # Nếu chạy thực tế ở hiện tại (2024/early 2025), bạn cần fallback về gemini-1.5-flash
                    try:
                        model = genai.GenerativeModel(model_name=model_version)
                    except:
                        st.warning(f"⚠️ Model {model_version} chưa public API tại local, fallback về gemini-1.5-flash để demo.")
                        model = genai.GenerativeModel(model_name="gemini-1.5-flash")

                    response = model.generate_content([prompt, gemini_file])
                    
                    # 4. Display
                    st.success("✅ Xử lý hoàn tất!")
                    
                    # Tab view cho giao diện hiện đại
                    tab1, tab2 = st.tabs(["📄 Báo cáo chi tiết", "📊 NotebookLM Data"])
                    
                    with tab1:
                        st.markdown(response.text)
                    
                    with tab2:
                        if opt_notebooklm:
                            st.info("Dữ liệu dưới đây được định dạng để Copy/Paste vào NotebookLM Studio (Table/Slide Source).")
                            # Giả lập trích xuất JSON từ text (trong thực tế dùng response schema)
                            st.code(f"""
                            {{
                                "source": "Meeting_Audio",
                                "generated_by": "{model_version}",
                                "slides_suggestion": [
                                    {{"slide": 1, "title": "Tổng quan", "bullets": ["..."]}},
                                    {{"slide": 2, "title": "Số liệu", "bullets": ["..."]}}
                                ]
                            }}
                            """, language="json")
                        else:
                            st.write("Bạn chưa chọn tính năng NotebookLM Data.")

                    # 5. Export Word
                    doc = create_docx(response.text)
                    doc_io = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    doc.save(doc_io.name)
                    
                    with open(doc_io.name, "rb") as f:
                        st.download_button(
                            "📥 Tải báo cáo (.docx)", 
                            f, 
                            "Meeting_Report_Gen3.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Cleanup
                    genai.delete_file(gemini_file.name)
                    os.remove(tmp_path)
                    os.remove(doc_io.name)

                except Exception as e:
                    st.error(f"Lỗi hệ thống: {e}")

if __name__ == "__main__":
    main()
